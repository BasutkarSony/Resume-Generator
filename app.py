import streamlit as st
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor

# Streamlit Page Config
st.set_page_config(page_title="SmartResume Generator", page_icon="📄")

# Function to create a professional, single-page resume
def create_resume(data, photo):
    doc = Document()

    # Set Font Style for formal look
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)

    # Create a table for profile picture and contact details
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    cell1, cell2 = table.rows[0].cells

    # Add Profile Photo (if available) in Round Shape
    if photo:
        run = cell1.paragraphs[0].add_run()
        run.add_picture(photo, width=Inches(1.5))
        cell1.width = Inches(1.5)

    # Add Name & Contact Information (Beside the Image)
    profile_text = f"\n{data['name']}\n📧 {data['email']}\n📞 {data['phone']}\n🌐 {data['linkedin']}"
    profile_paragraph = cell2.paragraphs[0]
    run = profile_paragraph.add_run(profile_text)
    run.bold = True
    profile_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Add Job Title in Bold & Dark Blue
    title = doc.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = title.add_run(data["job_title"])
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 0, 139)  # Dark Blue for emphasis

    # Profile Summary
    doc.add_heading("Profile", level=2)
    doc.add_paragraph(data["profile_summary"])

    # Education (Compact Format)
    doc.add_heading("Education", level=2)
    for edu in data["education"]:
        doc.add_paragraph(f" {edu}", style='List Bullet')

    # Skills (Bullet Points)
    doc.add_heading("Skills", level=2)
    for skill in data["skills"]:
        doc.add_paragraph(f" {skill}", style='List Bullet')

    # Work Experience (Compact)
    doc.add_heading("Professional Experience", level=2)
    for exp in data["experience"]:
        doc.add_paragraph(f" {exp}", style='List Bullet')

    # Projects (Compact)
    doc.add_heading("Projects", level=2)
    for proj in data["projects"]:
        doc.add_paragraph(f" {proj}", style='List Bullet')

    # Certifications (Compact)
    doc.add_heading("Certifications", level=2)
    for cert in data["certifications"]:
        doc.add_paragraph(f" {cert}", style='List Bullet')

    # Languages (Inline Format)
    doc.add_heading("Languages", level=2)
    doc.add_paragraph(", ".join(data["languages"]))

    # Save the document
    filename = "Generated_Resume.docx"
    doc.save(filename)
    return filename

# Streamlit UI
def main():
    st.title("📄 SmartResume Generator")

    # Sidebar for Profile Picture Upload
    st.sidebar.header("📸 Upload Profile Picture")
    photo = st.sidebar.file_uploader("Upload Image", type=["jpg", "png", "jpeg"])

    # Personal Information
    st.header("Personal Information")
    name = st.text_input("Full Name")
    job_title = st.text_input("Job Title")
    email = st.text_input("Email")
    phone = st.text_input("Phone Number")
    linkedin = st.text_input("LinkedIn URL")

    # Profile Summary
    st.header("Profile Summary")
    profile_summary = st.text_area("Write a short summary about yourself")

    # Education
    st.header("Education")
    education = st.text_area("Enter your education details").split(",")

    # Skills
    st.header("Skills")
    skills = st.text_area("Enter your skills (e.g., Python, Java, SQL)").split(",")

    # Experience
    st.header("Work Experience")
    experience = st.text_area("Enter your work experience").split(",")

    # Projects
    st.header("Projects")
    projects = st.text_area("Enter your projects").split(",")

    # Certifications
    st.header("Certifications")
    certifications = st.text_area("Enter your certifications ").split(",")

    # Languages
    st.header("Languages")
    languages = st.text_area("Enter the languages you know").split(",")

    # Generate Resume Button
    if st.button("✨ Generate Resume"):
        if name and job_title and email and phone and linkedin and profile_summary:
            # Store data
            resume_data = {
                "name": name,
                "job_title": job_title,
                "email": email,
                "phone": phone,
                "linkedin": linkedin,
                "profile_summary": profile_summary,
                "education": education,
                "skills": skills,
                "experience": experience,
                "projects": projects,
                "certifications": certifications,
                "languages": languages
            }

            # Generate Resume
            resume_file = create_resume(resume_data, photo)

            # Provide Download Button
            with open(resume_file, "rb") as file:
                st.download_button(
                    label="📥 Download Resume",
                    data=file,
                    file_name=resume_file,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
        else:
            st.error("⚠️ Please fill in all required fields.")

# Run the app
if __name__ == "__main__":
    main()
